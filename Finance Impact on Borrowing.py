from docx import Document

# Create a Word document
doc = Document()

# Title
doc.add_heading("Borrowing Scenarios Report", level=1)

# Introduction
doc.add_paragraph(
    "This report presents three borrowing capacity scenarios based on varying income levels, "
    "expenses, and interest rates. The scenarios are designed to assist stakeholders in assessing "
    "the impact of financial factors on loan affordability."
)

# Scenario 1
doc.add_heading("Scenario 1: Moderate Income, Low Expenses", level=2)
doc.add_paragraph(
    "Income: $4,200/month\n"
    "Expenses: $2,000/month\n"
    "Interest Rate: 5.0%\n"
    "Borrowing Capacity: $132,000 (over 5 years)"
)
doc.add_paragraph("This scenario assumes a moderate income with minimal expenses, resulting in a high borrowing capacity.")

# Scenario 2
doc.add_heading("Scenario 2: High Income, High Expenses", level=2)
doc.add_paragraph(
    "Income: $5,000/month\n"
    "Expenses: $3,500/month\n"
    "Interest Rate: 4.8%\n"
    "Borrowing Capacity: $90,000 (over 5 years)"
)
doc.add_paragraph("This scenario highlights the impact of high expenses, which significantly reduce borrowing capacity despite a higher income.")

# Scenario 3
doc.add_heading("Scenario 3: Low Income, Low Expenses", level=2)
doc.add_paragraph(
    "Income: $3,500/month\n"
    "Expenses: $1,800/month\n"
    "Interest Rate: 5.2%\n"
    "Borrowing Capacity: $101,000 (over 5 years)"
)
doc.add_paragraph("This scenario demonstrates the balance between low income and low expenses, resulting in a moderate borrowing capacity.")

# Conclusion
doc.add_heading("Conclusion", level=2)
doc.add_paragraph(
    "The scenarios above illustrate how varying financial factors impact borrowing capacity. "
    "These examples provide a foundation for decision-making and discussions with stakeholders."
)

# Save the document
file_path = "Borrowing_Scenarios_Report.docx"
doc.save(file_path)

print(f"Report saved as {file_path}")
